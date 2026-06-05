import os
import sys
import numpy as np
from docx import Document
from docx.shared import Inches
import base64
import io
from PIL import Image

# Ensure project root is on sys.path for cr_analysis import
project_root = r"k:/app/4. AHP마스터"
if project_root not in sys.path:
    sys.path.append(project_root)

from cr_analysis import run_analysis, generate_report, matrix_to_heatmap_img

def generate_random_matrix(n, consistency=0.8):
    rng = np.random.default_rng()
    a = rng.random((n, n))
    a = (a + a.T) / 2
    np.fill_diagonal(a, 1.0)
    a = consistency * a + (1 - consistency) * np.ones((n, n))
    return a

# Generate sample data
original = generate_random_matrix(4, consistency=0.85)
corrected = original * 1.02  # simulated correction
metrics = run_analysis(original, corrected, "Sample Correction")

# Create Word document
doc = Document()
doc.add_heading('CR Analysis Report', level=1)

# Add table with metrics
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Option'
hdr_cells[1].text = 'Euclidean'
hdr_cells[2].text = 'Manhattan'
hdr_cells[3].text = 'Cosine Sim.'
hdr_cells[4].text = 'Distortion Score'
row_cells = table.add_row().cells
row_cells[0].text = metrics['option']
row_cells[1].text = f"{metrics['euclidean']:.4f}"
row_cells[2].text = f"{metrics['manhattan']:.4f}"
row_cells[3].text = f"{metrics['cosine_similarity']:.4f}"
row_cells[4].text = f"{metrics['distortion_score']:.4f}"

# Add heatmap images
def add_heatmap_image(matrix, title):
    img_base64 = matrix_to_heatmap_img(matrix, title)
    img_bytes = base64.b64decode(img_base64)
    img = Image.open(io.BytesIO(img_bytes))
    img_path = os.path.join(project_root, f"{title.replace(' ', '_')}.png")
    img.save(img_path)
    doc.add_heading(title, level=2)
    doc.add_picture(img_path, width=Inches(5))
    os.remove(img_path)

add_heatmap_image(original, "Original Matrix")
add_heatmap_image(corrected, "Corrected Matrix (Sample Correction)")

# Save docx
output_path = os.path.join(project_root, "cr_analysis_report.docx")
doc.save(output_path)
print(f"Word report saved to {output_path}")
