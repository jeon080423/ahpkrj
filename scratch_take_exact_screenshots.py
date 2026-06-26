import os
import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

async def capture_precise_screenshots():
    output_dir = "AHP마스터 블로그"
    os.makedirs(output_dir, exist_ok=True)
    
    img1_path = os.path.join(output_dir, "screenshot1_exact.png")
    img2_path = os.path.join(output_dir, "screenshot2_exact.png")
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 1280, "height": 900})
        
        print("Connecting to local Streamlit app...")
        await page.goto("http://localhost:8501", timeout=60000)
        
        print("Waiting for Streamlit to render...")
        await page.wait_for_timeout(8000)
        
        # Click Tab 2
        print("Clicking Tab 2...")
        try:
            tabs = await page.locator('button[data-baseweb="tab"]').all()
            if len(tabs) >= 2:
                await tabs[1].click()
                await page.wait_for_timeout(4000)
        except Exception as e:
            print("Failed to click tab:", e)
            
        # Image 1: The top part with info messages
        print("Taking first screenshot (top of Tab 2)...")
        await page.screenshot(path=img1_path)
        
        # Image 2: Scroll to Section 2 or 3
        print("Scrolling to Section 2...")
        try:
            # We look for the text "섹션 2"
            section2 = page.get_by_text("섹션 2: AHP 모델 계층구조 설계")
            await section2.scroll_into_view_if_needed()
            await page.wait_for_timeout(2000)
            await page.screenshot(path=img2_path)
        except Exception as e:
            print("Failed to find section 2:", e)
            # fallback
            await page.mouse.wheel(0, 1500)
            await page.wait_for_timeout(2000)
            await page.screenshot(path=img2_path)
            
        await browser.close()
        return img1_path, img2_path

def create_docx(img1, img2):
    output_dir = "AHP마스터 블로그"
    doc_path = os.path.join(output_dir, "AHP 상대비료 온라인 설문 작성 및 배포 무료_최종수정본.docx")
    
    doc = Document()
    
    title = doc.add_heading("AHP 상대비료 온라인 설문 작성 및 배포 무료", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p1 = doc.add_paragraph()
    p1.add_run("최근 의사결정 계층화 분석법(AHP, Analytic Hierarchy Process)을 활용한 논문 작성 및 기업 내 의사결정 수요가 꾸준히 증가하고 있습니다. AHP 설문지는 일반적인 리커트 척도 설문과는 달리, 요소 간의 상대적인 중요도를 쌍대비교(Pairwise Comparison)하는 방식으로 구성되어야 하므로 직접 설문 폼을 제작하기가 매우 까다롭습니다.").bold = False
    
    p2 = doc.add_paragraph("하지만 이제 'AHP 마스터' 애플리케이션의 ")
    p2.add_run("온라인 설문 작성 및 배포 기능").bold = True
    p2.add_run("을 이용하면 이러한 고민을 완벽하게 해결할 수 있습니다. 본 글에서는 해당 기능의 유용성과 놀라운 편의성에 대해 상세히 안내해 드립니다.")
    
    doc.add_heading("1. 비회원도 누구나! 무료로 체험하는 설문지 자동 생성", level=2)
    doc.add_paragraph("가장 큰 장점 중 하나는 별도의 비용이나 복잡한 가입 절차 없이, 웹사이트에 접속하는 것만으로 누구나 AHP 설문지 폼을 직접 작성해보고 미리보기 기능까지 경험할 수 있다는 점입니다. 1계층과 2계층의 평가 지표만 입력하면, 시스템이 알아서 가능한 모든 조합의 쌍대비교 문항을 생성해 줍니다.")
    
    doc.add_heading("2. 구글 스프레드시트와의 강력한 실시간 연동", level=2)
    doc.add_paragraph("설문지 폼 작성을 마쳤다면, 버튼 클릭 한 번으로 배포용 URL이 생성됩니다. 이 URL을 카카오톡이나 이메일 등으로 응답자들에게 전달하기만 하면 됩니다. 응답자들이 설문을 제출하면, 그 결과가 내 구글 스프레드시트에 실시간으로 차곡차곡 쌓이게 되어 별도의 데이터 취합 과정이 전혀 필요하지 않습니다.")
    
    doc.add_heading("3. 단 한 번의 가입으로 누리는 '작성 내용 자동 보존' 기능", level=2)
    doc.add_paragraph("비회원 상태에서 정성껏 입력한 설문 문항들이 날아갈까 걱정하지 않으셔도 됩니다. 폼 작성 도중에 화면을 닫지 않고 바로 사이드바를 통해 무료 회원가입을 완료하면, 방금 전까지 입력했던 모든 내용이 마법처럼 그대로 유지됩니다. 가입 직후 즉시 [배포 및 DB 연동] 버튼을 눌러 설문을 시작할 수 있는 최적의 사용자 경험(UX)을 제공합니다.")
    
    doc.add_paragraph()
    
    doc.add_heading("실제 웹 화면 미리보기", level=2)
    doc.add_paragraph("아래는 실제 AHP 마스터 플랫폼에서 제공하는 온라인 설문지 작성 화면의 모습입니다. 직관적이고 깔끔한 인터페이스를 통해 누구나 손쉽게 설문 문항을 세팅할 수 있습니다.")
    
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(6.0))
        p = doc.add_paragraph("▲ 비회원도 접근 가능한 온라인 설문 작성 탭과 친절한 안내 화면")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(6.0))
        p = doc.add_paragraph("▲ 계층 지표 입력 시 자동으로 쌍대비교 문항이 구성되는 편리한 UI")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph("지금 바로 'AHP 마스터'에 접속하여 무료로 제공되는 막강한 설문 배포 및 데이터 수집 기능을 직접 경험해 보세요. 연구와 업무의 효율성이 비약적으로 향상될 것입니다!")
    doc.add_paragraph()
    p_url = doc.add_paragraph("🌐 AHP 마스터 바로가기: ")
    p_url.add_run("https://ahpkrj.streamlit.app/").bold = True
    doc.add_paragraph()
    doc.add_paragraph("#AHP #AHP마스터 #쌍대비교 #설문조사 #논문작성 #의사결정 #설문지자동생성 #무료설문지")
    
    doc.save(doc_path)
    print(f"Blog post created at {doc_path}")

async def main():
    img1, img2 = await capture_precise_screenshots()
    create_docx(img1, img2)

if __name__ == "__main__":
    asyncio.run(main())
