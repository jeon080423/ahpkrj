import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_new_blog():
    output_dir = "AHP마스터 블로그"
    os.makedirs(output_dir, exist_ok=True)
    
    img1 = r"C:\Users\jeon0\.gemini\antigravity-ide\brain\15892de2-4a9e-4137-ac6e-2dbd63dc7b59\media__1781818293830.png"
    img2 = r"C:\Users\jeon0\.gemini\antigravity-ide\brain\15892de2-4a9e-4137-ac6e-2dbd63dc7b59\media__1781818350839.png"
    
    doc_path = os.path.join(output_dir, "쌍대비교 척도 무료 사용_AHP마스터.docx")
    
    doc = Document()
    
    # 제목
    title = doc.add_heading("쌍대비교 척도 무료 사용: AHP 분석을 위해 특화된 설문 플랫폼이 필요한 이유", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 도입부
    p_intro = doc.add_paragraph()
    p_intro.add_run("의사결정 계층화 분석법(AHP, Analytic Hierarchy Process)은 다수의 대안 중 최적의 의사결정을 내리기 위해 널리 사용되는 분석 기법입니다. 이 분석을 위해서는 응답자가 두 요소를 1:1로 비교하는 ").bold = False
    p_intro.add_run("'쌍대비교(Pairwise Comparison)'").bold = True
    p_intro.add_run(" 척도가 필수적입니다.")
    
    # 1. 일반 폼의 한계
    doc.add_heading("1. 구글 폼과 네이버 폼에서는 쌍대비교 척도를 제공하지 않습니다", level=2)
    doc.add_paragraph("많은 연구자와 실무자들이 설문조사를 진행할 때 가장 익숙한 구글 폼(Google Forms)이나 네이버 폼을 먼저 떠올립니다. 하지만 이러한 일반적인 설문 폼은 5점 척도나 7점 척도와 같은 리커트(Likert) 척도는 잘 지원하지만, 좌우 양극단의 요소를 9점 척도로 비교해야 하는 '쌍대비교 척도' 형식은 제공하지 않습니다.")
    doc.add_paragraph("억지로 객관식 문항이나 그리드 형식을 변형하여 쌍대비교 폼을 만들 수는 있지만, 응답자가 보기에 매우 직관적이지 않고 응답 과정에서 엄청난 피로도를 유발하게 됩니다. 결국 이는 응답률 저하와 불성실 응답으로 이어지게 됩니다. 따라서 AHP 분석을 위해서는 쌍대비교에 특화된 전용 설문 서비스를 반드시 이용해야만 합니다.")
    
    # 2. AHP 마스터의 쌍대비교 척도 UI
    doc.add_heading("2. 직관적이고 편리한 AHP 마스터의 쌍대비교 척도 (무료 제공)", level=2)
    doc.add_paragraph("'AHP 마스터' 플랫폼은 AHP 분석을 위한 완벽한 쌍대비교 척도 UI를 누구나 무료로 사용할 수 있도록 제공합니다. 좌우 대칭형의 깔끔한 디자인을 통해 응답자가 어떤 요소를 더 중요하게 생각하는지 직관적으로 클릭하여 응답할 수 있습니다.")
    
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(6.0))
        p = doc.add_paragraph("▲ AHP 마스터에서 제공하는 실제 쌍대비교 척도 UI 화면")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[이미지 삽입 오류: 첫 번째 업로드 이미지를 찾을 수 없습니다]")
        
    doc.add_paragraph()
    
    # 3. 일관성 보조 기능
    doc.add_heading("3. 논리적 일관적인 응답을 유도하는 '일관성 가이드' 기능", level=2)
    doc.add_paragraph("쌍대비교 설문에서 가장 빈번하게 발생하는 문제는 바로 응답자의 '논리적 일관성(CR, Consistency Ratio)' 결여입니다. 예를 들어 A가 B보다 중요하고 B가 C보다 중요하다고 응답한 후, C가 A보다 중요하다고 응답하는 모순이 발생하면 해당 데이터는 폐기되어야 합니다.")
    doc.add_paragraph("'AHP 마스터'는 이러한 문제를 사전에 방지하기 위해, 응답자가 비논리적인 패턴으로 응답할 경우 실시간으로 이를 감지하고 일관성 있는 응답을 하도록 유도하는 '응답 일관성 가이드' 기능을 제공합니다. 이 혁신적인 기능을 통해 버려지는 데이터 없이 논문 통과를 보장하는 고품질의 설문 데이터를 확보할 수 있습니다.")
    
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(6.0))
        p = doc.add_paragraph("▲ 논리적 모순을 방지하고 일관성 있는 응답을 돕는 가이드 화면")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[이미지 삽입 오류: 두 번째 업로드 이미지를 찾을 수 없습니다]")
        
    doc.add_paragraph()
    doc.add_paragraph("이처럼 구글 폼이나 네이버 폼이 대체할 수 없는 AHP 전문 기능들을 무료로 제공하는 'AHP 마스터'를 통해 여러분의 연구와 실무 분석을 성공적으로 이끌어 보세요!")
    
    doc.add_paragraph()
    doc.add_paragraph("#AHP #AHP마스터 #쌍대비교 #쌍대비교척도 #설문조사 #구글폼 #네이버폼 #논문작성 #일관성비율 #CR가이드")
    
    doc.save(doc_path)
    print(f"Blog post created at {doc_path}")

if __name__ == "__main__":
    create_new_blog()
